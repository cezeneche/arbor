from __future__ import annotations

import os
import threading
from collections import defaultdict
from io import BytesIO
from typing import TYPE_CHECKING

from fastapi import HTTPException

if TYPE_CHECKING:
    from PIL import Image


def _tables_to_text(tables: list) -> str:
    lines = []
    for table in tables:
        if not isinstance(table, list):
            continue
        for row in table:
            if not isinstance(row, list):
                continue
            cells = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)

# PaddleOCR model init takes 30-60s on first load. Keep a single instance for
# the lifetime of the process so subsequent calls pay no init cost.
_paddle_ocr_instance: object = None
_paddle_ocr_lock = threading.Lock()


def _get_paddle_ocr() -> object:
    global _paddle_ocr_instance
    if _paddle_ocr_instance is not None:
        return _paddle_ocr_instance
    with _paddle_ocr_lock:
        if _paddle_ocr_instance is None:
            try:
                from paddleocr import PaddleOCR  # type: ignore
                try:
                    _paddle_ocr_instance = PaddleOCR(use_textline_orientation=True, lang="en")
                except TypeError:
                    _paddle_ocr_instance = PaddleOCR(use_angle_cls=True, lang="en")
            except Exception as exc:
                raise HTTPException(status_code=500, detail="paddleocr not installed") from exc
    return _paddle_ocr_instance


def _is_text(filename: str, content_type: str | None) -> bool:
    lower_name = filename.lower()
    return lower_name.endswith(".txt") or bool(content_type and content_type.startswith("text/"))


def _is_image(filename: str, content_type: str | None) -> bool:
    lower_name = filename.lower()
    return lower_name.endswith((".png", ".jpg", ".jpeg")) or bool(
        content_type and content_type.startswith("image/")
    )


def _is_pdf(filename: str, content_type: str | None) -> bool:
    lower_name = filename.lower()
    return lower_name.endswith(".pdf") or content_type == "application/pdf"


def _is_docx(filename: str, content_type: str | None) -> bool:
    lower_name = filename.lower()
    return lower_name.endswith(".docx") or content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )


def _extract_docx(data: bytes) -> dict[str, object]:
    try:
        from docx import Document as DocxDocument  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail="python-docx not installed") from exc
    try:
        doc = DocxDocument(BytesIO(data))
        parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return {"raw_text": "\n".join(parts), "pages": [], "layout": None}
    except Exception as exc:
        raise HTTPException(status_code=422, detail="The Word document could not be opened. Try saving it as a PDF and upload again.") from exc


def _decode_text_bytes(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")


def _to_sequence(value):
    if isinstance(value, (str, bytes)):
        return None
    if isinstance(value, (list, tuple)):
        return list(value)
    try:
        return list(value)
    except Exception:
        return None


def _bbox_to_xy(bbox) -> tuple[float, float] | None:
    seq = _to_sequence(bbox)
    if not seq:
        return None

    first = seq[0]
    first_seq = _to_sequence(first)
    if first_seq and len(first_seq) >= 2:
        try:
            return float(first_seq[0]), float(first_seq[1])
        except (TypeError, ValueError):
            return None

    if len(seq) >= 2:
        try:
            return float(seq[0]), float(seq[1])
        except (TypeError, ValueError):
            return None
    return None


def _to_float(value, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _bbox_to_rect(bbox) -> list[float] | None:
    seq = _to_sequence(bbox)
    if not seq:
        return None

    if len(seq) >= 4 and all(not isinstance(x, (list, tuple, dict)) for x in seq[:4]):
        x1 = _to_float(seq[0])
        y1 = _to_float(seq[1])
        x2 = _to_float(seq[2])
        y2 = _to_float(seq[3])
        return [x1, y1, x2, y2]

    points: list[tuple[float, float]] = []
    for item in seq:
        point = _to_sequence(item)
        if point and len(point) >= 2:
            points.append((_to_float(point[0]), _to_float(point[1])))

    if not points:
        if len(seq) >= 2:
            x = _to_float(seq[0])
            y = _to_float(seq[1])
            return [x, y, x, y]
        return None

    xs = [pt[0] for pt in points]
    ys = [pt[1] for pt in points]
    return [min(xs), min(ys), max(xs), max(ys)]


def _normalize_ocr_lines(result) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    seen: set[tuple[str, tuple[float, float, float, float]]] = set()

    def add_line(text, bbox=None, confidence=None) -> None:
        if text is None:
            return
        text_value = str(text).strip()
        if not text_value:
            return
        rect = _bbox_to_rect(bbox) or [0.0, 0.0, 0.0, 0.0]
        key = (text_value, (rect[0], rect[1], rect[2], rect[3]))
        if key in seen:
            return
        seen.add(key)
        rows.append(
            {
                "text": text_value,
                "bbox": rect,
                "confidence": _to_float(confidence),
            }
        )

    def walk(node) -> None:
        if node is None:
            return

        if isinstance(node, dict):
            rec_texts = node.get("rec_texts")
            if isinstance(rec_texts, (list, tuple)):
                poly_candidates = (
                    node.get("dt_polys")
                    or node.get("rec_polys")
                    or node.get("dt_boxes")
                    or node.get("rec_boxes")
                    or []
                )
                polys = _to_sequence(poly_candidates) or []
                for idx, text_item in enumerate(rec_texts):
                    bbox_item = polys[idx] if idx < len(polys) else None
                    score_item = None
                    rec_scores = node.get("rec_scores")
                    if isinstance(rec_scores, (list, tuple)) and idx < len(rec_scores):
                        score_item = rec_scores[idx]
                    add_line(text_item, bbox_item, score_item)

            text = node.get("rec_text") or node.get("text") or node.get("label")
            bbox = (
                node.get("dt_polys")
                or node.get("dt_boxes")
                or node.get("bbox")
                or node.get("box")
                or node.get("points")
            )
            confidence = node.get("score") or node.get("confidence")
            if text is not None:
                add_line(text, bbox, confidence)

            for value in node.values():
                if isinstance(value, (list, tuple, dict)):
                    walk(value)
            return

        if isinstance(node, (list, tuple)):
            if len(node) >= 2:
                bbox = node[0]
                text_info = node[1]
                text = None
                confidence = None
                if isinstance(text_info, (list, tuple)) and text_info:
                    text = text_info[0]
                    if len(text_info) > 1:
                        confidence = text_info[1]
                elif isinstance(text_info, str):
                    text = text_info
                if text is not None:
                    add_line(text, bbox, confidence)
                    return

            for item in node:
                if isinstance(item, (list, tuple, dict)):
                    walk(item)

    walk(result)
    rows.sort(key=lambda row: (_to_float(row["bbox"][1]), _to_float(row["bbox"][0])))
    return rows


def _build_layout_blocks(
    ocr_lines: list[dict[str, object]],
    image_height: int,
) -> dict[str, list[dict[str, object]]]:
    grouped: dict[str, list[int]] = defaultdict(list)
    safe_height = max(image_height, 1)

    for idx, line in enumerate(ocr_lines):
        bbox = line.get("bbox")
        rect = bbox if isinstance(bbox, list) and len(bbox) == 4 else [0.0, 0.0, 0.0, 0.0]
        y_center = (_to_float(rect[1]) + _to_float(rect[3])) / 2.0
        ratio = y_center / float(safe_height)
        if ratio <= 0.25:
            grouped["header"].append(idx)
        elif ratio >= 0.80:
            grouped["footer"].append(idx)
        else:
            grouped["body"].append(idx)

    blocks: list[dict[str, object]] = []
    for block_type in ("header", "body", "footer"):
        idxs = grouped.get(block_type, [])
        if not idxs:
            continue
        text = " ".join(str(ocr_lines[i].get("text", "")).strip() for i in idxs).strip()
        blocks.append({"type": block_type, "text": text, "lines_idx": idxs})
    return {"blocks": blocks}


def _extract_image_document_with_paddleocr(data: bytes) -> dict[str, object]:
    try:
        import cv2  # type: ignore
        import numpy as np  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail="paddleocr not installed") from exc

    arr = np.frombuffer(data, dtype=np.uint8)
    image = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if image is None:
        raise HTTPException(status_code=422, detail="The image could not be opened. Check the file is a valid PNG or JPG.")
    image_height = int(image.shape[0]) if getattr(image, "shape", None) is not None else 1

    ocr = _get_paddle_ocr()

    try:
        result = ocr.predict(image)
    except TypeError:
        # older versions
        result = ocr.ocr(image)

    ocr_lines = _normalize_ocr_lines(result)
    raw_text = "\n".join(str(line["text"]) for line in ocr_lines).strip()
    return {
        "raw_text": raw_text,
        "ocr_lines": ocr_lines,
        "layout": _build_layout_blocks(ocr_lines, image_height=image_height),
    }


def _extract_text_with_paddleocr(data: bytes) -> str:
    return str(_extract_image_document_with_paddleocr(data).get("raw_text", ""))


def _extract_text_with_paddleocr_image(image: "Image.Image") -> str:
    try:
        import cv2  # type: ignore
        import numpy as np  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail="paddleocr not installed") from exc

    image_rgb = image.convert("RGB")
    image_np = np.array(image_rgb)
    image_bgr = cv2.cvtColor(image_np, cv2.COLOR_RGB2BGR)

    ocr = _get_paddle_ocr()

    try:
        result = ocr.predict(image_bgr)
    except TypeError:
        # older versions
        result = ocr.ocr(image_bgr)

    ocr_lines = _normalize_ocr_lines(result)
    return "\n".join(str(line["text"]) for line in ocr_lines).strip()


def _extract_pdf_document_hybrid(data: bytes) -> dict[str, object]:
    try:
        import pdfplumber  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail="pdfplumber not installed") from exc

    pages_text: list[dict[str, object]] = []
    try:
        with pdfplumber.open(BytesIO(data)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                text = (page.extract_text() or "").strip()
                table_text = ""
                try:
                    tables = page.extract_tables() or []
                    table_text = _tables_to_text(tables)
                except Exception:
                    table_text = ""
                combined_page_text = "\n".join(filter(None, [text, table_text]))
                words_payload: list[dict[str, object]] = []
                try:
                    words = page.extract_words() or []
                    if isinstance(words, list):
                        for word in words:
                            if not isinstance(word, dict):
                                continue
                            words_payload.append(
                                {
                                    "text": str(word.get("text", "")).strip(),
                                    "x0": _to_float(word.get("x0"), 0.0),
                                    "y0": _to_float(word.get("top"), 0.0),
                                    "x1": _to_float(word.get("x1"), 0.0),
                                    "y1": _to_float(word.get("bottom"), 0.0),
                                }
                            )
                except Exception:
                    words_payload = []
                pages_text.append(
                    {
                        "page_number": page_idx,
                        "text": combined_page_text,
                        "source": "pdf_text",
                        "words": words_payload,
                    }
                )
    except Exception as exc:
        raise HTTPException(status_code=422, detail="The PDF could not be opened. Try re-saving or re-exporting it and upload again.") from exc

    combined_text = "\n\n".join(str(page["text"]) for page in pages_text if str(page["text"]).strip()).strip()
    if len(combined_text) >= 100:
        return {
            "raw_text": combined_text,
            "pages": pages_text,
            "layout": None,
        }

    try:
        from pdf2image import convert_from_bytes  # type: ignore
    except Exception as exc:
        raise HTTPException(status_code=500, detail="pdf2image not installed") from exc

    try:
        images = convert_from_bytes(data)
    except Exception as exc:
        raise HTTPException(status_code=422, detail="The scanned PDF could not be processed. Try converting it to a standard PDF and upload again.") from exc

    pages_ocr: list[dict[str, object]] = []
    for page_idx, image in enumerate(images[:3], start=1):  # cap at 3 pages — goods lines can be on page 2
        page_text = _extract_text_with_paddleocr_image(image).strip()
        pages_ocr.append(
            {
                "page_number": page_idx,
                "text": page_text,
                "source": "pdf_ocr",
            }
        )

    combined_ocr_text = "\n\n".join(str(page["text"]) for page in pages_ocr if str(page["text"]).strip()).strip()
    if not combined_ocr_text:
        raise HTTPException(status_code=422, detail="No text could be read from this PDF. If it is a scanned document, ensure the scan is clear and try again.")

    return {
        "raw_text": combined_ocr_text,
        "pages": pages_ocr,
        "layout": None,
    }


def extract_document_from_upload(filename: str, content_type: str | None, data: bytes) -> dict[str, object]:
    if _is_text(filename, content_type):
        raw_text = _decode_text_bytes(data)
        return {
            "raw_text": raw_text,
            "ocr_lines": [],
            "layout": {"blocks": []},
        }

    if _is_pdf(filename, content_type):
        return _extract_pdf_document_hybrid(data)

    if _is_docx(filename, content_type):
        return _extract_docx(data)

    if _is_image(filename, content_type):
        if os.getenv("OCR_DISABLED") == "1":
            return {"raw_text": "", "ocr_lines": [], "layout": {"blocks": []}}
        return _extract_image_document_with_paddleocr(data)

    raise HTTPException(status_code=415, detail="This file type is not supported. Upload a PDF, Word document, PNG, or JPG.")


def extract_text_from_upload(filename: str, content_type: str | None, data: bytes) -> str:
    extracted = extract_document_from_upload(
        filename=filename,
        content_type=content_type,
        data=data,
    )
    return str(extracted.get("raw_text", ""))
